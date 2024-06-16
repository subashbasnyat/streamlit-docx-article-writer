import io
import os
import streamlit as st
from docx import Document
from typing import List
from loguru import logger
from langchain_community.chat_models import ChatOllama
from langchain_nvidia_ai_endpoints import ChatNVIDIA
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import PromptTemplate
from langchain_core.pydantic_v1 import BaseModel, Field


os.environ["NVIDIA_API_KEY"] = "XXXXXXXXXXXXXX"

llm_ollama = ChatOllama(model="llama3", temperature=0.8, top_p=0.8)
llm_nvidia = ChatNVIDIA(
    model="mistralai/mixtral-8x7b-instruct-v0.1", temperature=0.8, top_p=0.8
)

llm = llm_nvidia


class DocumentModel(BaseModel):
    filename: str = Field(description="relevant name of the docx file based on heading")
    heading: str = Field(description="main heading of the document")
    paragraphs: List[str] = Field(description="5 paragraphs about the heading")


def create_docx_files(filename, heading, paragraphs):
    # Create a new Document
    doc = Document()
    # Add a title
    doc.add_heading(f"{heading}", level=1)

    # Add some sample content
    doc.add_paragraph(f"{paragraphs}")

    # Save the document
    if filename.endswith(".docx"):
        doc_filename = f"{filename}"
    else:
        doc_filename = f"{filename}.docx"
    return doc, doc_filename


def generate_document(keyword, writing_style, word_count):
    document_query = f"""Write me a document about {keyword}.
    The document should be {word_count} words long.
    The document should be in {writing_style} style.
    The filename should should not contain extra characters.
    """

    parser = JsonOutputParser(pydantic_object=DocumentModel)

    prompt = PromptTemplate(
        template="Answer the user query.\n{format_instructions}\n{query}\n",
        input_variables=["query"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )

    chain = prompt | llm | parser

    document_content = chain.invoke({"query": document_query})
    if document_content.get("properties") is not None:
        document_content = document_content.get("properties")
        file_property = document_content.get("filename")
        filename = file_property.get("title")
        heading_property = document_content.get("heading")
        heading = heading_property.get("title")
        paragraphs = document_content.get("paragraphs")
    else:
        filename = document_content.get("filename")
        heading = document_content.get("heading")
        paragraphs = document_content.get("paragraphs")
    if type(paragraphs) is list:
        paragraphs = "\n".join(paragraphs)
    logger.info(f"""Filename: {filename}, Heading: {heading}""")
    logger.debug(f"""Paragraphs: {paragraphs}""")
    logger.debug(f"""Document Content: {document_content}""")
    return filename, heading, paragraphs


st.title("Article Writer")
st.subheader("Write an article about any topic")

keyword = st.text_input("Enter a keyword")
writing_style = st.selectbox(
    "Select a writing style", ["Academic", "Business", "Sarcastic", "Casual", "Funny"]
)
word_count = st.slider(
    "Select a word count", min_value=300, max_value=1000, step=100, value=300
)

submit_buttton = st.button("Generate Article")

if submit_buttton:
    message = st.empty()
    message.text("Generating article...")
    filename, heading, paragraphs = generate_document(
        keyword, writing_style, word_count
    )
    message.text("")
    st.write(paragraphs)

    doc_download, doc_filename = create_docx_files(filename, heading, paragraphs)

    bio = io.BytesIO()
    doc_download.save(bio)
    if doc_download:
        st.download_button(
            label="Click here to download (docx)",
            data=bio.getvalue(),
            file_name=f"{doc_filename}",
            mime="docx",
        )
